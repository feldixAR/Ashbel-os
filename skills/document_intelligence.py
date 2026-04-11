"""
skills/document_intelligence.py — Document Intelligence Skill
Phase 4: Document upload → structured leads/entities.

Stateless. No DB. No AI calls (unless explicitly requested).
Accepts file bytes + metadata, returns normalized records for
feeding into existing LeadOps / lead_intelligence pipeline.

Supported formats:
  CSV   — stdlib csv (always available)
  Excel — openpyxl (optional, graceful fallback)
  Word  — python-docx (optional, graceful fallback)
  PDF   — pdfplumber (optional, graceful fallback)
  TXT   — stdlib

All output records are normalized to the same shape as
skills.lead_intelligence.normalize() input.
"""

from __future__ import annotations
import csv
import io
import logging
from dataclasses import dataclass, field
from typing import Any

log = logging.getLogger(__name__)

# ── Result contract ───────────────────────────────────────────────────────────

@dataclass
class ParsedDocument:
    file_name:    str
    format:       str               # "csv" | "excel" | "word" | "pdf" | "text" | "unknown"
    row_count:    int               # number of data rows found
    records:      list[dict]        # normalized lead-like records
    raw_tables:   list[list[list]]  # [[row, ...], ...] — raw detected tables
    warnings:     list[str]         = field(default_factory=list)
    text_blocks:  list[str]         = field(default_factory=list)   # free-text paragraphs


# ── Main entry ────────────────────────────────────────────────────────────────

def parse_document(content: bytes, file_name: str = "",
                   hint_format: str = "") -> ParsedDocument:
    """
    Parse document bytes into structured records.
    hint_format: "csv" | "excel" | "word" | "pdf" | "text" — optional override.
    """
    fmt = hint_format or _detect_format(file_name, content)

    try:
        if fmt == "csv":
            return _parse_csv(content, file_name)
        if fmt == "excel":
            return _parse_excel(content, file_name)
        if fmt == "word":
            return _parse_word(content, file_name)
        if fmt == "pdf":
            return _parse_pdf(content, file_name)
        if fmt == "text":
            return _parse_text(content, file_name)
    except Exception as e:
        log.warning(f"[DocIntel] parse failed format={fmt}: {e}")
        return ParsedDocument(file_name=file_name, format=fmt,
                              row_count=0, records=[], raw_tables=[],
                              warnings=[f"parse error: {e}"])

    return ParsedDocument(file_name=file_name, format="unknown",
                          row_count=0, records=[], raw_tables=[],
                          warnings=["unsupported format"])


def detect_lead_columns(headers: list[str]) -> dict:
    """
    Map column headers to lead fields.
    Returns dict: {field: column_index}
    """
    mapping = {}
    for i, h in enumerate(headers):
        hl = (h or "").lower().strip()
        if any(w in hl for w in ["שם", "name", "full name", "שם מלא", "contact"]):
            mapping.setdefault("name", i)
        elif any(w in hl for w in ["טלפון", "phone", "mobile", "נייד", "tel"]):
            mapping.setdefault("phone", i)
        elif any(w in hl for w in ["מייל", "email", "mail", "דוא\"ל"]):
            mapping.setdefault("email", i)
        elif any(w in hl for w in ["עיר", "city", "ישוב", "location"]):
            mapping.setdefault("city", i)
        elif any(w in hl for w in ["חברה", "company", "עסק", "business", "firm"]):
            mapping.setdefault("company", i)
        elif any(w in hl for w in ["תפקיד", "role", "job", "position", "direct"]):
            mapping.setdefault("role", i)
        elif any(w in hl for w in ["הערות", "notes", "פרטים", "info", "description"]):
            mapping.setdefault("notes", i)
    return mapping


def normalize_records(rows: list[list], headers: list[str],
                      source_type: str = "document") -> list[dict]:
    """Convert raw table rows into lead-normalized dicts."""
    col_map = detect_lead_columns(headers)
    records = []
    for row in rows:
        if not any(cell for cell in row):
            continue   # skip empty rows
        rec = {
            "name":        _cell(row, col_map.get("name")),
            "phone":       _cell(row, col_map.get("phone")),
            "email":       _cell(row, col_map.get("email")),
            "city":        _cell(row, col_map.get("city")),
            "company":     _cell(row, col_map.get("company")),
            "role":        _cell(row, col_map.get("role")),
            "notes":       _cell(row, col_map.get("notes")),
            "source_type": source_type,
        }
        # Skip rows with no identifying information
        if rec["name"] or rec["phone"] or rec["email"]:
            records.append(rec)
    return records


# ── Format parsers ────────────────────────────────────────────────────────────

def _parse_csv(content: bytes, file_name: str) -> ParsedDocument:
    text = content.decode("utf-8-sig", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        return ParsedDocument(file_name=file_name, format="csv",
                              row_count=0, records=[], raw_tables=[])
    headers = rows[0]
    data    = rows[1:]
    records = normalize_records(data, headers, source_type="csv_import")
    return ParsedDocument(
        file_name=file_name, format="csv",
        row_count=len(data), records=records,
        raw_tables=[rows],
    )


def _parse_excel(content: bytes, file_name: str) -> ParsedDocument:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
        ws = wb.active
        rows = [[str(cell.value or "") for cell in row] for row in ws.iter_rows()]
        if not rows:
            return ParsedDocument(file_name=file_name, format="excel",
                                  row_count=0, records=[], raw_tables=[])
        headers = rows[0]
        data    = rows[1:]
        records = normalize_records(data, headers, source_type="excel_import")
        return ParsedDocument(
            file_name=file_name, format="excel",
            row_count=len(data), records=records,
            raw_tables=[rows],
        )
    except ImportError:
        # Fallback: try parsing as CSV (xlsx without openpyxl)
        log.warning("[DocIntel] openpyxl not available — trying CSV fallback")
        try:
            return _parse_csv(content, file_name)
        except Exception:
            return ParsedDocument(file_name=file_name, format="excel",
                                  row_count=0, records=[], raw_tables=[],
                                  warnings=["openpyxl not installed"])


def _parse_word(content: bytes, file_name: str) -> ParsedDocument:
    try:
        import docx
        doc = docx.Document(io.BytesIO(content))
        text_blocks = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_raw  = []
        all_records = []
        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            tables_raw.append(rows)
            if rows:
                records = normalize_records(rows[1:], rows[0], source_type="word_import")
                all_records.extend(records)
        combined_text = "\n".join(text_blocks)
        # If no table records, try extracting leads from free text
        if not all_records and combined_text:
            all_records = _extract_from_freetext(combined_text)
        return ParsedDocument(
            file_name=file_name, format="word",
            row_count=len(all_records), records=all_records,
            raw_tables=tables_raw, text_blocks=text_blocks,
        )
    except ImportError:
        log.warning("[DocIntel] python-docx not available")
        return ParsedDocument(file_name=file_name, format="word",
                              row_count=0, records=[], raw_tables=[],
                              warnings=["python-docx not installed"])


def _parse_pdf(content: bytes, file_name: str) -> ParsedDocument:
    try:
        import pdfplumber
        text_blocks = []
        tables_raw  = []
        all_records = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                txt = page.extract_text() or ""
                if txt.strip():
                    text_blocks.append(txt)
                for table in (page.extract_tables() or []):
                    if not table:
                        continue
                    tables_raw.append(table)
                    records = normalize_records(table[1:], table[0] or [],
                                               source_type="pdf_import")
                    all_records.extend(records)
        combined = "\n".join(text_blocks)
        if not all_records and combined:
            all_records = _extract_from_freetext(combined)
        return ParsedDocument(
            file_name=file_name, format="pdf",
            row_count=len(all_records), records=all_records,
            raw_tables=tables_raw, text_blocks=text_blocks,
        )
    except ImportError:
        log.warning("[DocIntel] pdfplumber not available")
        return ParsedDocument(file_name=file_name, format="pdf",
                              row_count=0, records=[], raw_tables=[],
                              warnings=["pdfplumber not installed"])


def _parse_text(content: bytes, file_name: str) -> ParsedDocument:
    text = content.decode("utf-8", errors="replace")
    records = _extract_from_freetext(text)
    return ParsedDocument(
        file_name=file_name, format="text",
        row_count=len(records), records=records,
        raw_tables=[], text_blocks=[text],
    )


# ── Freetext lead extraction ──────────────────────────────────────────────────

def _extract_from_freetext(text: str) -> list[dict]:
    """Extract lead-like entities from unstructured text using regex."""
    import re
    records = []
    phone_re = re.compile(r"(\+?0?[0-9]{8,13})")
    email_re = re.compile(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z]{2,}")

    phones = phone_re.findall(text)
    emails = email_re.findall(text)

    if phones or emails:
        records.append({
            "name":        "",
            "phone":       phones[0] if phones else "",
            "email":       emails[0] if emails else "",
            "city":        "",
            "company":     "",
            "role":        "",
            "notes":       text[:300],
            "source_type": "text_extract",
        })
    return records


# ── Helpers ───────────────────────────────────────────────────────────────────

def _detect_format(file_name: str, content: bytes) -> str:
    fn = (file_name or "").lower()
    if fn.endswith(".csv"):                      return "csv"
    if fn.endswith((".xlsx", ".xls")):           return "excel"
    if fn.endswith((".docx", ".doc")):           return "word"
    if fn.endswith(".pdf"):                      return "pdf"
    if fn.endswith(".txt"):                      return "text"
    # Magic bytes
    if content[:4] == b"PK\x03\x04":            return "excel"   # ZIP-based (xlsx/docx)
    if content[:4] == b"%PDF":                  return "pdf"
    return "text"


def _cell(row: list, idx: int | None) -> str:
    if idx is None or idx >= len(row):
        return ""
    return str(row[idx] or "").strip()
